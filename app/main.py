import streamlit as st
import os
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import re  # For cleaning text

# Additional imports for alternative OCR
import fitz  # PyMuPDF
import easyocr
from PIL import Image
import numpy as np

# Load environment variables
load_dotenv()

client = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key=os.environ.get("GROQ_API_KEY")
)

# Set page configuration with a modern layout
st.set_page_config(
    page_title="AI Wash Service",
    page_icon="🤖",
    layout="wide"
)

# Apply custom CSS for a modern look
st.markdown("""
    <style>
        /* General page styling */
        body {
            background-color: #f0f2f6;
            color: #333;
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        }
        /* Title and subtitle styling */
        h1, h3 {
            color: #2c3e50;
            text-align: center;
            font-weight: 600;
        }
        /* Paragraph styling */
        p {
            color: #34495e;
            text-align: center;
            font-size: 1.1rem;
        }
        /* File uploader styling */
        .stFileUploader label {
            color: #2980b9;
            font-weight: bold;
        }
        /* Button styling */
        .stButton button {
            background-color: #3498db;
            color: white;
            border-radius: 5px;
            padding: 10px 20px;
            font-size: 1rem;
            transition: background-color 0.3s ease;
        }
        .stButton button:hover {
            background-color: #2980b9;
        }
        /* Text area styling */
        .stTextArea textarea {
            border-radius: 5px;
            border: 1px solid #bdc3c7;
        }
        /* Footer styling */
        footer {
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9rem;
        }
    </style>
""", unsafe_allow_html=True)

# Title and description with a modern UI
st.title("🚀 AI Wash Service")
st.markdown("""
    <p>
        This app helps you process and refine your documents using AI. Upload a file and let AI enhance its content!
    </p>
""", unsafe_allow_html=True)

# File uploader for TXT, DOCX, PDF, and Images
st.markdown("<h3>📤 Upload Your Document</h3>", unsafe_allow_html=True)
uploaded_file = st.file_uploader("", type=['txt', 'pdf', 'docx', 'png', 'jpeg', 'jpg'])

def extract_text_from_pdf(file):
    """
    Extract text from a PDF file. If standard extraction yields very little text,
    use PyMuPDF to render each page and EasyOCR to extract text.
    """
    # Read file bytes for OCR if needed
    file_bytes = file.read()
    file.seek(0)  # Reset file pointer for PdfReader
    reader = PdfReader(file)
    extracted_text = "\n".join([page.extract_text() or "" for page in reader.pages]).strip()
    
    # If extracted text is very short, assume OCR is needed
    if len(extracted_text) < 50:
        try:
            # Open PDF with PyMuPDF from the byte stream
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            ocr_text = ""
            # Initialize EasyOCR (using English)
            reader_ocr = easyocr.Reader(['en'], gpu=False)
            for page in doc:
                pix = page.get_pixmap()
                # Convert the pixmap to a PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                # Convert image to numpy array
                img_np = np.array(img)
                # Extract text using EasyOCR
                result = reader_ocr.readtext(img_np, detail=0, paragraph=True)
                ocr_text += " ".join(result) + "\n"
            return ocr_text.strip()
        except Exception as e:
            st.error(f"EasyOCR extraction failed: {str(e)}")
            return ""
    return extracted_text

def extract_text_from_docx(file):
    """Extract text from a DOCX file."""
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_text_from_image(file):
    """
    Extract text from an image file using EasyOCR.
    """
    try:
        img = Image.open(file)
        # Convert image to RGB and then to numpy array
        img_np = np.array(img.convert("RGB"))
        reader_ocr = easyocr.Reader(['en'], gpu=False)
        result = reader_ocr.readtext(img_np, detail=0, paragraph=True)
        return " ".join(result).strip()
    except Exception as e:
        st.error(f"Image OCR extraction failed: {str(e)}")
        return ""

content = ""

if uploaded_file is not None:
    st.markdown(f"<p>📄 Uploaded file: <b>{uploaded_file.name}</b></p>", unsafe_allow_html=True)
    
    if uploaded_file.type == "text/plain":
        content = uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        content = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        content = extract_text_from_docx(uploaded_file)
    elif uploaded_file.type in ["image/png", "image/jpeg", "image/jpg"]:
        content = extract_text_from_image(uploaded_file)
    else:
        st.warning("⚠ Unsupported file type.")
    
    if content:
        with st.expander("📄 Preview Extracted File Content", expanded=True):
            st.text_area("File Content Preview", content, height=250, disabled=True)

if "final_content" not in st.session_state:
    st.session_state.final_content = ""

if content and st.session_state.final_content == "":
    st.markdown("---")
    st.markdown("<h3>✨ AI-Enhanced Content</h3>", unsafe_allow_html=True)

    system_prompt = """You are an exceptional writer with expertise in crafting formal, informative, and engaging articles and papers. Your task is to refine and enhance the provided content, ensuring it is clear, well-structured, and compelling while maintaining a professional and authoritative tone. Improve coherence, depth, and readability, remove redundancies, vague statements, or filler content, and focus on delivering factually accurate and insightful writing.
Don't mention that you made changes. Just provide the refined article content."""
    try:
        llm_response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content}
            ]
        )
        enhanced_content = llm_response.choices[0].message.content
        if enhanced_content:
            cleaned_content = re.sub(r"\*", "", enhanced_content)
            st.session_state.final_content = cleaned_content
    except Exception as e:
        st.error(f"Error processing AI response: {str(e)}")

final_content = st.text_area("✏️ AI-Enhanced Content (Editable)", st.session_state.final_content, height=250, key="final_edit")

extra_instructions = st.text_area("📝 Extra Instructions (Optional)", 
                                  "Enter any additional instructions to modify the content...", 
                                  height=100)
if extra_instructions and st.button("Apply Extra Instructions"):
    st.markdown("---")
    st.markdown("<h3>🔄 Updating Content with Extra Instructions</h3>", unsafe_allow_html=True)
    update_prompt = f"""You are a refined writer tasked with updating the following content according to the extra instructions provided.
    
Extra Instructions: {extra_instructions}

Update the content below accordingly. Do not mention that changes were made.

Content:
{final_content}
    """
    try:
        update_response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": update_prompt},
                {"role": "user", "content": final_content}
            ]
        )
        updated_content = update_response.choices[0].message.content
        if updated_content:
            cleaned_updated_content = re.sub(r"\*", "", updated_content)
            st.session_state.final_content = cleaned_updated_content
            final_content = cleaned_updated_content
    except Exception as e:
        st.error(f"Error updating AI response: {str(e)}")

def generate_docx(text):
    """Generates a clean DOCX file from the given text without a title."""
    doc = Document()
    paragraphs = text.split("\n\n")
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if final_content:
    st.markdown("---")
    st.markdown("<h3>📥 Download AI-Enhanced Content</h3>", unsafe_allow_html=True)
    cleaned_docx = generate_docx(re.sub(r"\*", "", final_content))
    st.download_button(
        label="📄 Download as DOCX",
        data=cleaned_docx,
        file_name="Enhanced_Content.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.markdown("---")
st.markdown(
    "<footer>Celly Services (CSI) 2025</footer>",
    unsafe_allow_html=True
)
